from typing import Dict, List, Optional
import re
import docx2txt
import PyPDF2
from docx import Document
import numpy as np
from collections import Counter

class DocumentFeatureExtractor:
    """Extract document-specific features for improved classification"""
    
    def __init__(self):
        # Common document section keywords
        self.section_keywords = {
            'academic_credentials': ['academic', 'education', 'qualification', 'degree', 'diploma'],
            'certification': ['certification', 'certified', 'certificate'],
            'transcript': ['transcript', 'grades', 'marks', 'subjects'],
            'service_record': ['service', 'employment', 'work', 'position', 'designation']
        }
        
        # Document structure patterns
        self.structure_patterns = {
            'academic_credentials': [
                r'degree|diploma|certificate of|awarded to',
                r'university|college|institution',
                r'graduation|completed|conferred'
            ],
            'certification': [
                r'this is to certify|hereby certifies?',
                r'earned units?|completed units?',
                r'course|program|training'
            ],
            'transcript': [
                r'transcript of records?',
                r'academic record',
                r'grade[s]?|mark[s]?|score[s]?'
            ],
            'service_record': [
                r'service record|employment record',
                r'position|designation',
                r'date of appointment|period of service'
            ]
        }

    def extract_features(self, text: str, file_path: str = None) -> Dict:
        """Extract comprehensive document features"""
        features = {}
        
        # Basic text features
        features.update(self._extract_text_features(text))
        
        # Layout and formatting features
        if file_path:
            if file_path.lower().endswith('.docx'):
                features.update(self._extract_docx_features(file_path))
            elif file_path.lower().endswith('.pdf'):
                features.update(self._extract_pdf_features(file_path))
        
        # Structural features
        features.update(self._extract_structural_features(text))
        
        # Section-based features
        features.update(self._extract_section_features(text))
        
        return features

    def _extract_text_features(self, text: str) -> Dict:
        """Extract basic text features"""
        features = {}
        
        # Text length features
        features['text_length'] = len(text)
        features['word_count'] = len(text.split())
        features['avg_word_length'] = np.mean([len(w) for w in text.split()])
        
        # Line features
        lines = text.split('\n')
        features['line_count'] = len(lines)
        features['avg_line_length'] = np.mean([len(line) for line in lines])
        
        # Character type ratios
        total_chars = len(text)
        features['uppercase_ratio'] = sum(1 for c in text if c.isupper()) / total_chars if total_chars > 0 else 0
        features['digit_ratio'] = sum(1 for c in text if c.isdigit()) / total_chars if total_chars > 0 else 0
        features['punctuation_ratio'] = sum(1 for c in text if c in '.,;:!?-()[]{}') / total_chars if total_chars > 0 else 0
        
        return features

    def _extract_docx_features(self, file_path: str) -> Dict:
        """Extract DOCX-specific formatting features"""
        features = {}
        try:
            doc = Document(file_path)
            
            # Paragraph formatting
            para_formats = []
            for para in doc.paragraphs:
                if para.style:
                    para_formats.append(para.style.name)
            
            features['para_styles_count'] = len(set(para_formats))
            features['total_paragraphs'] = len(doc.paragraphs)
            
            # Table features
            features['table_count'] = len(doc.tables)
            
            # Font variations
            fonts = set()
            for p in doc.paragraphs:
                for run in p.runs:
                    if run.font:
                        fonts.add(run.font.name)
            features['font_variation_count'] = len(fonts)
            
        except Exception as e:
            print(f"Error extracting DOCX features: {str(e)}")
            features['docx_features_error'] = True
            
        return features

    def _extract_pdf_features(self, file_path: str) -> Dict:
        """Extract PDF-specific features"""
        features = {}
        try:
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                features['page_count'] = len(pdf.pages)
                
                # Extract text layout features
                text_blocks = []
                for page in pdf.pages:
                    text = page.extract_text()
                    blocks = text.split('\n\n')
                    text_blocks.extend(blocks)
                
                features['text_block_count'] = len(text_blocks)
                features['avg_block_length'] = np.mean([len(block) for block in text_blocks])
                
        except Exception as e:
            print(f"Error extracting PDF features: {str(e)}")
            features['pdf_features_error'] = True
            
        return features

    def _extract_structural_features(self, text: str) -> Dict:
        """Extract document structure features"""
        features = {}
        
        # Check for common document parts
        features['has_header'] = bool(re.search(r'^[^\n]{1,100}$', text.split('\n')[0]))
        features['has_footer'] = bool(re.search(r'^[^\n]{1,100}$', text.split('\n')[-1]))
        features['has_date'] = bool(re.search(r'\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{2,4}', text))
        
        # Section detection
        sections = text.split('\n\n')
        features['section_count'] = len(sections)
        
        # Table-like structure detection
        table_like_lines = sum(1 for line in text.split('\n') if line.count('\t') > 2 or line.count('  ') > 3)
        features['table_like_structure_ratio'] = table_like_lines / features['section_count'] if features['section_count'] > 0 else 0
        
        return features

    def _extract_section_features(self, text: str) -> Dict:
        """Extract features based on document sections and keywords"""
        features = {}
        
        # Check for document type-specific patterns
        for doc_type, patterns in self.structure_patterns.items():
            matches = 0
            for pattern in patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    matches += 1
            features[f'{doc_type}_pattern_matches'] = matches
            features[f'{doc_type}_pattern_ratio'] = matches / len(patterns)
        
        # Keyword presence
        for section, keywords in self.section_keywords.items():
            keyword_count = sum(1 for keyword in keywords if keyword in text.lower())
            features[f'{section}_keyword_count'] = keyword_count
            features[f'{section}_keyword_ratio'] = keyword_count / len(keywords)
        
        return features

    def normalize_features(self, features: Dict) -> Dict:
        """Normalize numerical features to a common scale"""
        normalized = {}
        for key, value in features.items():
            if isinstance(value, (int, float)):
                # Min-max normalization for numeric values
                if key.endswith('_ratio'):
                    # Ratios are already between 0 and 1
                    normalized[key] = value
                else:
                    # Normalize other numeric values using log scale
                    normalized[key] = np.log1p(value) if value > 0 else 0
            else:
                normalized[key] = value
        return normalized