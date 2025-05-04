import requests
import os
from pathlib import Path
import docx
import time
import glob

def create_sample_docx(filename="test_doc.docx"):
    """Create a sample Word document for testing"""
    doc = docx.Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document created for processing.')
    doc.add_paragraph('It contains different types of text to test extraction:')
    doc.add_paragraph('• Headers and paragraphs\n• Lists and formatting\n• Multiple sections')
    doc.save(filename)
    return filename

def test_document_processing(file_path):
    """Test document processing with a specific file"""
    url = 'http://127.0.0.1:8000/api/documents/process/'
    
    print(f"\nTesting document: {os.path.basename(file_path)}")
    print(f"File path: {file_path}")
    
    if not os.path.exists(file_path):
        print(f"Error: File {file_path} not found")
        return False
    
    try:
        with open(file_path, 'rb') as f:
            files = {'file': (os.path.basename(file_path), f)}
            print("\nSending request...")
            print(f"File size: {os.path.getsize(file_path)} bytes")
            
            response = requests.post(url, files=files)
            print('\nResponse Status Code:', response.status_code)
            
            if response.ok:
                data = response.json()
                print('\nExtracted Text:')
                print('-' * 60)
                print(data.get('extracted_text', 'No text extracted'))
                print('-' * 60)
                print('\nClassification:', data.get('classification', 'No classification'))
                print('Document ID:', data.get('document_id'))
                print('File Type:', data.get('file_type', 'Not specified'))
                return True
            else:
                print('\nError Response:')
                print(response.text)
                return False
            
    except Exception as e:
        print('Error:', e)
        return False

def run_tests():
    """Run tests with different document types"""
    base_dir = Path(__file__).resolve().parent
    test_docs_dir = base_dir / 'test_docs'
    
    print("\nLooking for test documents in:", test_docs_dir)
    
    # Find all PDF and DOCX files in the test_docs directory
    pdf_files = list(test_docs_dir.glob('*.pdf'))
    docx_files = list(test_docs_dir.glob('*.docx'))
    
    # Also test with our sample image
    test_files = [
        *pdf_files,  # All PDF files
        *docx_files,  # All DOCX files
        str(base_dir / '../assets/images/udmaddress.png')  # Test image
    ]
    
    if not pdf_files and not docx_files:
        print("\nNo PDF or DOCX files found in test_docs directory.")
        print("Please add some test documents to:", test_docs_dir)
        print("Supported formats: .pdf, .docx")
    
    print(f"\nFound {len(pdf_files)} PDF files and {len(docx_files)} DOCX files")
    print("\nStarting document processing tests...")
    successful_tests = 0
    
    for file_path in test_files:
        print("\n" + "="*80)
        if test_document_processing(str(file_path)):
            successful_tests += 1
        time.sleep(1)
    
    print("\n" + "="*80)
    print(f"\nTests completed. Successfully processed {successful_tests}/{len(test_files)} files.")

if __name__ == '__main__':
    run_tests()